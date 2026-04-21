"""Markdown and DOCX exports with provenance appendix."""

from __future__ import annotations

import json
from io import BytesIO
from typing import Any, Dict, List, Mapping, Sequence

import yaml

from llm_survey.prompts.model_extraction_prompts import EXTRACTION_SYSTEM_PROMPT


def _coverage(gap_report: Mapping[str, Any] | None) -> float:
    if not gap_report:
        return 0.0
    return float(
        gap_report.get("structural_coverage_score", gap_report.get("overall_model_completeness", 0.0)) or 0.0
    )


def _plain(payload: Any) -> Any:
    if hasattr(payload, "model_dump"):
        return payload.model_dump()
    return payload


def _validation_map(validations: Any) -> Dict[str, Dict[str, Any]]:
    plain = _plain(validations) or {}
    rows = plain.get("validations", []) if isinstance(plain, dict) else []
    output: Dict[str, Dict[str, Any]] = {}
    for row in rows:
        if not isinstance(row, dict):
            continue
        key = str(row.get("hypothesis_id", "")).strip()
        if key:
            output[key] = row
    return output


def build_methods_markdown(
    extraction_results: Sequence[Dict[str, Any]],
    gap_report: Mapping[str, Any] | None,
    chunk_lookup: Mapping[str, str],
) -> str:
    lines: List[str] = [
        "# Structured model draft (machine-assisted)",
        "",
        "## Structural coverage (heuristic)",
        f"- Score: **{_coverage(gap_report):.2f}** (see project docs for what this does and does not mean.)",
        f"- Testability (heuristic): **{float((gap_report or {}).get('model_testability_score', 0) or 0):.2f}**",
        "",
        "## Extracted content",
        "",
    ]
    ev_idx = 1
    evidence_lines: List[str] = []

    for row in extraction_results:
        if not row.get("success") or not row.get("model"):
            continue
        cid = str(row.get("chunk_id", ""))
        model = row["model"]
        lines.append(f"### Chunk `{cid}`")
        chunk_body = chunk_lookup.get(cid, "")
        for var in model.get("variables") or []:
            if not isinstance(var, dict):
                continue
            ev = var.get("evidence_strength", "direct")
            lines.append(
                f"- **Variable** ({ev}): **{var.get('name')}** — {var.get('definition', '')} "
                f"— quote: _{var.get('example_quote', '')}_"
            )
        for rel in model.get("relationships") or []:
            if not isinstance(rel, dict):
                continue
            ev = rel.get("evidence_strength", "direct")
            fq = str(rel.get("from_variable", ""))
            tq = str(rel.get("to_variable", ""))
            quote = str(rel.get("supporting_quote", ""))
            ids = rel.get("source_chunk_ids") or [cid]
            lines.append(
                f"- **Relationship** ({ev}): {fq} → {tq} — _{rel.get('mechanism', '')}_ "
                f"(confidence {rel.get('confidence', '')}) [evidence {ev_idx}]"
            )
            ctx = chunk_body
            snippet = quote if quote and quote in ctx else (ctx[:800] + "…" if len(ctx) > 800 else ctx)
            evidence_lines.append(f"{ev_idx}. Chunk(s) `{', '.join(str(i) for i in ids)}`: {snippet}")
            ev_idx += 1
        lines.append("")

    lines.append("## Evidence appendix")
    lines.extend(evidence_lines or ["_No relationship rows were exported._"])
    lines.append("")
    return "\n".join(lines)


def build_docx_bytes(
    extraction_results: Sequence[Dict[str, Any]],
    gap_report: Mapping[str, Any] | None,
    chunk_lookup: Mapping[str, str],
) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError as err:  # pragma: no cover
        raise ImportError("python-docx is required for DOCX export.") from err

    doc = Document()
    doc.add_heading("Structured model draft", level=0)
    p = doc.add_paragraph()
    p.add_run("Structural coverage (heuristic): ").bold = True
    p.add_run(f"{_coverage(gap_report):.2f}")
    doc.add_paragraph(
        "This document was generated to support a methods section draft. "
        "Verify every claim against the source quotations in the appendix."
    )

    doc.add_heading("Variables and relationships", level=1)
    n = 1
    appendix: List[tuple[int, str, str, str]] = []

    for row in extraction_results:
        if not row.get("success") or not row.get("model"):
            continue
        cid = str(row.get("chunk_id", ""))
        chunk_body = chunk_lookup.get(cid, "")
        for var in row["model"].get("variables") or []:
            if not isinstance(var, dict):
                continue
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"[{n}] ").bold = True
            para.add_run(
                f"Variable ({var.get('evidence_strength', 'direct')}): {var.get('name')} — {var.get('definition', '')}"
            )
            quote = str(var.get("example_quote", ""))
            ids = ", ".join(str(i) for i in (var.get("source_chunk_ids") or [cid]))
            snippet = quote if quote and quote in chunk_body else (chunk_body[:1200])
            appendix.append((n, ids, quote, snippet))
            n += 1
        for rel in row["model"].get("relationships") or []:
            if not isinstance(rel, dict):
                continue
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"[{n}] ").bold = True
            para.add_run(
                f"{rel.get('from_variable')} → {rel.get('to_variable')} "
                f"({rel.get('direction')}, {rel.get('evidence_strength', 'direct')}): "
            )
            para.add_run(str(rel.get("mechanism", "")))
            quote = str(rel.get("supporting_quote", ""))
            ids = ", ".join(str(i) for i in (rel.get("source_chunk_ids") or [cid]))
            snippet = quote if quote and quote in chunk_body else (chunk_body[:1200])
            appendix.append((n, ids, quote, snippet))
            n += 1

    doc.add_page_break()
    doc.add_heading("Evidence appendix", level=1)
    for num, ids, quote, snippet in appendix:
        h = doc.add_heading(f"Evidence {num}", level=2)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph(f"Source chunk id(s): {ids}")
        if quote:
            doc.add_paragraph("Quoted span:")
            doc.add_paragraph(quote, style="Intense Quote")
        doc.add_paragraph("Surrounding chunk context (truncated):")
        doc.add_paragraph(snippet)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_json_export_bundle(
    extraction_results: Sequence[Dict[str, Any]],
    gap_report: Mapping[str, Any] | None,
    chunk_lookup: Mapping[str, str],
    failure_summary: Mapping[str, Any] | None,
) -> str:
    return json.dumps(
        {
            "extraction_results": list(extraction_results),
            "gap_report": dict(gap_report or {}),
            "chunk_text_by_id": dict(chunk_lookup),
            "failure_summary": dict(failure_summary or {}),
            "system_prompt_sha_note": "SHA not computed; system prompt is versioned in repo.",
            "system_prompt_length": len(EXTRACTION_SYSTEM_PROMPT),
        },
        indent=2,
        ensure_ascii=False,
    )


def build_final_model_spec_yaml(
    consolidated_model: Mapping[str, Any] | Any,
    validations: Mapping[str, Any] | Any = None,
    conflict_report: Mapping[str, Any] | Any = None,
    metadata: Mapping[str, Any] | None = None,
) -> str:
    model = _plain(consolidated_model) or {}
    validation_by_id = _validation_map(validations)
    conflicts = (_plain(conflict_report) or {}).get("contradictions", [])
    meta = dict(metadata or {})

    payload = {
        "model": {
            "generated_at": meta.get("generated_at", ""),
            "pipeline_version": meta.get("pipeline_version", "1.0.0"),
            "total_chunks": meta.get("total_chunks", 0),
            "iterations_completed": meta.get("iterations_completed", 0),
            "overall_confidence": model.get("overall_confidence", 0.0),
            "model_summary": model.get("model_summary", ""),
        },
        "variables": [
            {
                "name": row.get("name"),
                "type": row.get("type"),
                "definition": row.get("definition"),
                "aliases": row.get("aliases", []),
                "chunk_frequency": row.get("chunk_frequency", 0),
                "confidence": row.get("confidence", 0.0),
            }
            for row in model.get("variables", [])
        ],
        "relationships": [
            {
                "from": row.get("from_variable"),
                "to": row.get("to_variable"),
                "direction": row.get("direction"),
                "mechanism": row.get("mechanism"),
                "confidence": row.get("confidence", 0.0),
                "support_count": row.get("support_count", 0),
                "support_fraction": row.get("support_fraction", 0.0),
            }
            for row in model.get("relationships", [])
        ],
        "hypotheses": [],
        "moderators": model.get("moderators", []),
        "contradictions": conflicts,
        "research_questions": model.get("research_questions", []),
    }

    for row in model.get("hypotheses", []):
        validation = validation_by_id.get(str(row.get("id", "")), {})
        payload["hypotheses"].append(
            {
                "id": row.get("id"),
                "statement": row.get("statement"),
                "confidence": row.get("confidence", 0.0),
                "support_count": row.get("support_count", 0),
                "support_fraction": row.get("support_fraction", 0.0),
                "consensus_strength": validation.get("consensus_strength", row.get("consensus_strength", "weak")),
                "literature_support_score": validation.get("literature_support_score", row.get("literature_support_score", 0.0)),
                "novelty_flag": validation.get("novelty_flag", row.get("novelty_flag", False)),
                "supporting_papers": validation.get("supporting_papers", []),
                "contradicting_papers": validation.get("contradicting_papers", []),
                "researcher_notes": row.get("researcher_notes", ""),
            }
        )

    return yaml.safe_dump(payload, sort_keys=False, allow_unicode=True)


def build_mermaid_diagram(consolidated_model: Mapping[str, Any] | Any) -> str:
    model = _plain(consolidated_model) or {}
    lines = ["graph LR"]
    for row in model.get("relationships", []):
        from_name = str(row.get("from_variable", "")).replace(" ", "")
        to_name = str(row.get("to_variable", "")).replace(" ", "")
        if not from_name or not to_name:
            continue
        label = f"{row.get('direction', 'unclear')}, conf:{float(row.get('confidence', 0.0)):.2f}"
        lines.append(f'    {from_name} -->|"{label}"| {to_name}')
    return "\n".join(lines)


def build_causal_graph_html(
    consolidated_model: Mapping[str, Any] | Any,
    validations: Mapping[str, Any] | Any = None,
    conflict_report: Mapping[str, Any] | Any = None,
) -> str:
    model = _plain(consolidated_model) or {}
    validation_by_id = _validation_map(validations)
    conflicts = (_plain(conflict_report) or {}).get("contradictions", [])
    mermaid = build_mermaid_diagram(model)

    relationship_cards: List[str] = []
    for row in model.get("relationships", []):
        quotes = "".join(f"<li>{quote}</li>" for quote in row.get("supporting_quotes", [])[:4])
        relationship_cards.append(
            "<details>"
            f"<summary><strong>{row.get('from_variable')}</strong> → <strong>{row.get('to_variable')}</strong> "
            f"({row.get('direction')}, conf {float(row.get('confidence', 0.0)):.2f})</summary>"
            f"<p>{row.get('mechanism', '')}</p>"
            f"<p>Support: {row.get('support_count', 0)} chunks ({float(row.get('support_fraction', 0.0)):.2f})</p>"
            f"<ul>{quotes}</ul>"
            "</details>"
        )

    hypothesis_cards: List[str] = []
    for row in model.get("hypotheses", []):
        validation = validation_by_id.get(str(row.get("id", "")), {})
        hypothesis_cards.append(
            "<details>"
            f"<summary>{row.get('id')}: {row.get('statement')}</summary>"
            f"<p>Confidence: {float(row.get('confidence', 0.0)):.2f}</p>"
            f"<p>Literature: {validation.get('consensus_strength', row.get('consensus_strength', 'weak'))} "
            f"(score {float(validation.get('literature_support_score', row.get('literature_support_score', 0.0))):.2f})</p>"
            "</details>"
        )

    conflict_cards = "".join(
        "<li>"
        f"{row.get('relationship')}: {row.get('resolution_status')} — {row.get('resolution_explanation')}"
        "</li>"
        for row in conflicts
    ) or "<li>No contradictions detected.</li>"

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Consolidated Causal Graph</title>
  <script src="https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.min.js"></script>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 24px; color: #18212b; line-height: 1.5; }}
    .grid {{ display: grid; grid-template-columns: 1.3fr 1fr; gap: 24px; }}
    .panel {{ border: 1px solid #d4dce6; border-radius: 12px; padding: 16px; background: #fafcfe; }}
    details {{ margin: 0 0 12px 0; }}
    summary {{ cursor: pointer; }}
  </style>
</head>
<body>
  <h1>Consolidated Causal Graph</h1>
  <p>{model.get("model_summary", "")}</p>
  <div class="grid">
    <section class="panel">
      <div class="mermaid">
{mermaid}
      </div>
    </section>
    <section class="panel">
      <h2>Contradictions</h2>
      <ul>{conflict_cards}</ul>
    </section>
  </div>
  <section class="panel">
    <h2>Relationships</h2>
    {''.join(relationship_cards)}
  </section>
  <section class="panel">
    <h2>Hypotheses</h2>
    {''.join(hypothesis_cards)}
  </section>
  <script>mermaid.initialize({{ startOnLoad: true }});</script>
</body>
</html>"""


def build_evidence_report_markdown(
    consolidated_model: Mapping[str, Any] | Any,
    validations: Mapping[str, Any] | Any = None,
    conflict_report: Mapping[str, Any] | Any = None,
) -> str:
    model = _plain(consolidated_model) or {}
    validation_by_id = _validation_map(validations)
    conflicts = (_plain(conflict_report) or {}).get("contradictions", [])
    lines: List[str] = [
        "# Evidence report",
        "",
        "## Model summary",
        model.get("model_summary", "_No summary available._"),
        "",
        "## Hypotheses",
        "",
    ]

    for row in model.get("hypotheses", []):
        validation = validation_by_id.get(str(row.get("id", "")), {})
        lines.append(f"### {row.get('id')}: {row.get('statement')}")
        lines.append(f"- Data confidence: **{float(row.get('confidence', 0.0)):.2f}**")
        lines.append(
            f"- Literature consensus: **{validation.get('consensus_strength', row.get('consensus_strength', 'weak'))}** "
            f"(score {float(validation.get('literature_support_score', row.get('literature_support_score', 0.0))):.2f})"
        )
        lines.append(f"- Novelty flag: **{bool(validation.get('novelty_flag', row.get('novelty_flag', False)))}**")
        for quote in row.get("supporting_quotes", [])[:5]:
            lines.append(f"- Supporting quote: _{quote}_")
        for paper in validation.get("supporting_papers", [])[:3]:
            lines.append(
                f"- Supporting paper: **{paper.get('title')}** ({paper.get('year', 'n.d.')}) — {paper.get('relevant_excerpt', '')}"
            )
        for paper in validation.get("contradicting_papers", [])[:3]:
            lines.append(
                f"- Contradicting paper: **{paper.get('title')}** ({paper.get('year', 'n.d.')}) — {paper.get('relevant_excerpt', '')}"
            )
        if row.get("researcher_notes"):
            lines.append(f"- Researcher notes: {row.get('researcher_notes')}")
        lines.append("")

    lines.extend(["## Contradictions", ""])
    if conflicts:
        for row in conflicts:
            lines.append(f"- **{row.get('relationship')}**: {row.get('resolution_status')} — {row.get('resolution_explanation')}")
    else:
        lines.append("- No contradictions detected.")
    lines.append("")
    return "\n".join(lines)
